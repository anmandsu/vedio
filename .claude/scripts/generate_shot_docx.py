#!/usr/bin/env python3
"""Generate DOCX shooting script for 针锋·针心 — 50 shots."""
import json, os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_docx(shots, out_path):
    doc = Document()
    for s in doc.sections:
        s.orientation = 1
        s.page_width = Cm(42)
        s.page_height = Cm(29.7)
        s.left_margin = Cm(1)
        s.right_margin = Cm(1)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('《针锋·针心》分镜头脚本')
    r.bold = True; r.font.size = Pt(18)
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.add_run('VPipe v2 · AI-Native | 是枝裕和体系 | 50镜').font.size = Pt(10)

    table = doc.add_table(rows=1, cols=10)
    table.style = 'Table Grid'
    for i, h in enumerate(['镜号','场','景别','机位/角度','焦段','运镜','光线','时长','画面内容','声音/备注']):
        c = table.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(8)

    for rd in shots:
        row = table.add_row()
        for i, txt in enumerate(rd):
            row.cells[i].text = txt
            for p in row.cells[i].paragraphs:
                for r in p.runs: r.font.size = Pt(7)

    doc.save(out_path)
    print(f'Saved: {out_path}')

if __name__ == '__main__':
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    json_path = os.path.join(base, 'projects', 'zhenfeng', 'shots', 'shot_data.json')
    out = os.path.join(base, 'projects', 'zhenfeng', 'shots', '针锋针心_分镜头脚本_50镜.docx')
    shots = json.loads(open(json_path, encoding='utf-8').read())
    build_docx(shots, out)
