#!/usr/bin/env python3
"""Convert VPipe YAML shot scripts to DOCX format."""

import yaml
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
from pathlib import Path


def add_heading(doc, text, level=1):
    """Add a heading with custom formatting."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_shot_to_doc(doc, shot, shot_number):
    """Add a single shot to the document."""
    # Shot header
    heading = doc.add_heading(f"镜头 {shot_number}: {shot['shot_id']}", level=2)

    # Shot metadata table
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light Grid Accent 1'

    metadata = [
        ("场景ID", shot.get('scene_id', '')),
        ("角色", shot.get('character', '')),
        ("地点", shot.get('location', '')),
        ("时长", f"{shot.get('duration_sec', 0)}秒"),
        ("节奏类型", shot.get('beat_type', '')),
    ]

    for label, value in metadata:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(value)

    doc.add_paragraph()

    # Shot details
    details = [
        ("剧情", 'plot'),
        ("表演", 'performance'),
        ("光线", 'lighting'),
        ("机位", 'camera'),
        ("调度", 'blocking'),
        ("声音", 'sound'),
        ("对白", 'dialogue'),
        ("风格注释", 'style_notes'),
    ]

    for label, key in details:
        if key in shot and shot[key]:
            p = doc.add_paragraph()
            p.add_run(f"{label}：").bold = True
            p.add_run(shot[key].strip())

    # Visual anchors
    if 'visual_anchors' in shot and shot['visual_anchors']:
        p = doc.add_paragraph()
        p.add_run("视觉锚点：").bold = True
        for anchor in shot['visual_anchors']:
            doc.add_paragraph(f"• {anchor}", style='List Bullet')

    # Continuity
    continuity = []
    if 'continuity_from' in shot:
        continuity.append(f"接续：{shot['continuity_from']}")
    if 'continuity_to' in shot:
        continuity.append(f"连接：{shot['continuity_to']}")

    if continuity:
        p = doc.add_paragraph()
        p.add_run("连续性：").bold = True
        p.add_run(" | ".join(continuity))

    # Generation notes
    if 'generation_notes' in shot and shot['generation_notes']:
        p = doc.add_paragraph()
        p.add_run("生成注释：").bold = True
        p.add_run(shot['generation_notes'].strip())
        run = p.runs[-1]
        run.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_paragraph("─" * 50)


def convert_yaml_to_docx(yaml_path, output_path):
    """Convert a VPipe YAML file to DOCX."""
    with open(yaml_path, 'r', encoding='utf-8') as f:
        data = yaml.safe_load(f)

    doc = Document()

    # Title
    title = doc.add_heading(f"第{data['episode']}集：{data['title']}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Style description
    p = doc.add_paragraph()
    p.add_run("风格：").bold = True
    p.add_run(data['style'])
    doc.add_paragraph()

    # Shots
    for idx, shot in enumerate(data['shots'], 1):
        add_shot_to_doc(doc, shot, idx)

    # Summary
    doc.add_page_break()
    doc.add_heading("统计信息", level=1)

    total_duration = sum(shot.get('duration_sec', 0) for shot in data['shots'])

    summary = [
        f"总镜头数：{len(data['shots'])}个",
        f"总时长：{total_duration}秒（{total_duration // 60}分{total_duration % 60}秒）",
    ]

    for line in summary:
        doc.add_paragraph(line, style='List Bullet')

    doc.save(output_path)
    print(f"[OK] 已生成：{output_path}")


def main():
    project_dir = Path(__file__).parent.parent.parent / "projects" / "daguangsai"
    shots_dir = project_dir / "shots"
    output_dir = project_dir / "docs"
    output_dir.mkdir(exist_ok=True)

    episodes = [
        ("episode_1.yaml", "第1集_拆牌倒计时_分镜头脚本.docx"),
        ("episode_2.yaml", "第2集_福袋崩裂_分镜头脚本.docx"),
        ("episode_3.yaml", "第3集_一针新生_分镜头脚本.docx"),
    ]

    for yaml_file, docx_file in episodes:
        yaml_path = shots_dir / yaml_file
        output_path = output_dir / docx_file

        if yaml_path.exists():
            convert_yaml_to_docx(yaml_path, output_path)
        else:
            print(f"[ERROR] 文件不存在：{yaml_path}")

    print(f"\n所有文档已生成到：{output_dir}")


if __name__ == "__main__":
    main()
